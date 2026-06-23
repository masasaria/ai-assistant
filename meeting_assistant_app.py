#!/usr/bin/env python3
"""
Meeting Assistant v2 — Gemma 4 edition
Run: python app.py  →  http://localhost:5000

Architecture:
  Upload → Gemma 4 processes full document (KB + concept map) → Start Listening unlocks
  Question → local concept-map expansion → keyword retrieval → Gemma 4 answers from short context

  Processing is one-time per session (~1-2 min for large docs).
  Answers are fast (~5-8s) because context is small (retrieval, not full doc).
  No concurrent Gemma 4 jobs — processing fully completes before any question is answered.

Requirements:
  pip install flask ollama faster-whisper python-docx pdfplumber
  ollama pull gemma4:12b && ollama serve
  (SpeechRecognition kept as a fallback — optional)
"""

import base64
import io
import json
import numpy as np
import re
import threading
import wave
import webbrowser
from pathlib import Path

from flask import Flask, jsonify, request, Response

try:
    import ollama as _ollama
except ImportError:
    print("Missing: pip install ollama"); raise

try:
    from docx import Document as DocxDocument
except ImportError:
    print("Missing: pip install python-docx"); raise

try:
    import pdfplumber
except ImportError:
    print("Missing: pip install pdfplumber"); raise

try:
    import anthropic as _anthropic
    _ANTHROPIC_AVAILABLE = True
except ImportError:
    _ANTHROPIC_AVAILABLE = False
    print("  [Anthropic] not installed — run: pip install anthropic")

try:
    import fitz as _fitz   # pymupdf
    _FITZ_AVAILABLE = True
    print("  [PDF] pymupdf ready — image extraction enabled")
except ImportError:
    _FITZ_AVAILABLE = False
    print("  [PDF] pymupdf not installed — PDF images skipped")
    print("        To install: pip install pymupdf")

try:
    from faster_whisper import WhisperModel as _WhisperModel
    _WHISPER_AVAILABLE = True
    print("  [Whisper] faster-whisper available — loading 'small' model at startup…")
    # Load default model at startup in a background thread so Flask starts immediately
    def _preload_whisper():
        try:
            _get_whisper("small")
            print("  [Whisper] 'small' model ready")
        except Exception as e:
            print(f"  [Whisper] Preload failed: {e}")
except ImportError:
    _WHISPER_AVAILABLE = False
    print("  [Whisper] Not installed — falling back to Google Speech")
    print("            To install: pip install faster-whisper")

try:
    import speech_recognition as sr
    _sr = sr.Recognizer()
    _SR_AVAILABLE = True
except ImportError:
    _SR_AVAILABLE = False

app = Flask(__name__)

DEFAULT_MODEL = "gemma4:12b"

_store = {
    # document
    "doc_text":    "",
    "doc_name":    "",
    "paragraphs":  [],   # raw first, KB appended after processing
    "raw_count":   0,    # boundary: paragraphs[:raw_count] = raw, rest = KB
    "index":       {},
    "images":      [],
    "tag_map":     {},   # citation tag → "Author (Year)" — from Word citation XML
    # knowledge base
    "summary":     "",
    "concept_map": {},
    # state
    "processing":  False,
    "proc_error":  "",
    "model":       DEFAULT_MODEL,
    "provider":    "ollama",
    "api_key":     "",
}


# ── Ollama ─────────────────────────────────────────────────────────────────────

def _to_ollama_messages(messages: list) -> list:
    """
    Convert OpenAI-style content blocks to the format the ollama library expects.
    ollama.Message requires content to be a plain string; images go in a separate
    'images' list as raw base64 strings (no data-URL prefix).
    """
    result = []
    for msg in messages:
        content = msg["content"]
        if isinstance(content, str):
            result.append({"role": msg["role"], "content": content})
        else:
            # content is a list of blocks (text + image_url / image)
            text_parts = []
            images     = []
            for block in content:
                btype = block.get("type", "")
                if btype == "text":
                    text_parts.append(block["text"])
                elif btype == "image_url":
                    url = block["image_url"]["url"]
                    # strip "data:image/xxx;base64," prefix
                    if "," in url:
                        images.append(url.split(",", 1)[1])
                elif btype == "image":
                    src = block.get("source", {})
                    if src.get("type") == "base64":
                        images.append(src["data"])
                # input_audio blocks are silently skipped (not supported yet)
            entry = {"role": msg["role"], "content": "\n".join(text_parts)}
            if images:
                entry["images"] = images
            result.append(entry)
    return result


def _to_anthropic_messages(messages: list) -> list:
    result = []
    for msg in messages:
        content = msg["content"]
        if isinstance(content, str):
            result.append({"role": msg["role"], "content": content})
        else:
            blocks = []
            for block in content:
                btype = block.get("type", "")
                if btype == "text":
                    blocks.append({"type": "text", "text": block["text"]})
                elif btype == "image_url":
                    url = block["image_url"]["url"]
                    if "," in url:
                        header, data = url.split(",", 1)
                        mt = header.split(":")[1].split(";")[0]
                        blocks.append({"type": "image", "source": {
                            "type": "base64", "media_type": mt, "data": data
                        }})
                elif btype == "image":
                    isrc = block.get("source", {})
                    if isrc.get("type") == "base64":
                        blocks.append({"type": "image", "source": isrc})
            result.append({"role": msg["role"], "content": blocks})
    return result


def _chat_ollama(messages: list, max_tokens: int = 1000, num_ctx: int = 8192, system: str = "") -> str:
    model = _store["model"]
    use_think = "gemma4" in model
    ollama_msgs = _to_ollama_messages(messages)
    if system:
        ollama_msgs = [{"role": "system", "content": system}] + ollama_msgs
    resp = _ollama.chat(
        model=model,
        messages=ollama_msgs,
        **({"think": False} if use_think else {}),
        options={"num_predict": max_tokens, "num_ctx": num_ctx, "temperature": 0},
    )
    text = (resp.message.content or "").strip()
    print(f"  [Chat/Ollama] {len(text)} chars (model={model}, ctx={num_ctx})")
    return text


def _chat_ollama_kb(messages: list, max_tokens: int = 1000, num_ctx: int = 8192, system: str = "") -> str:
    """Call Ollama for KB building — uses ollama_model, bypasses provider setting."""
    model = _store.get("ollama_model", "gemma4:12b")
    use_think = "gemma4" in model
    ollama_msgs = _to_ollama_messages(messages)
    if system:
        ollama_msgs = [{"role": "system", "content": system}] + ollama_msgs
    resp = _ollama.chat(
        model=model,
        messages=ollama_msgs,
        **({"think": False} if use_think else {}),
        options={"num_predict": max_tokens, "num_ctx": num_ctx, "temperature": 0},
    )
    text = (resp.message.content or "").strip()
    print(f"  [Chat/OllamaKB] {len(text)} chars (model={model})")
    return text


def _stream_ollama(messages: list, max_tokens: int = 400, num_ctx: int = 8192, system: str = ""):
    model = _store["model"]
    use_think = "gemma4" in model
    ollama_msgs = _to_ollama_messages(messages)
    if system:
        ollama_msgs = [{"role": "system", "content": system}] + ollama_msgs
    stream = _ollama.chat(model=model, messages=ollama_msgs, stream=True,
        **({"think": False} if use_think else {}),
        options={"num_predict": max_tokens, "num_ctx": num_ctx, "temperature": 0,
                 "keep_alive": -1})
    for chunk in stream:
        token = (chunk.message.content or "")
        if token:
            yield token


def _stream_ollama_kb(messages: list, max_tokens: int = 400, num_ctx: int = 8192, system: str = ""):
    """Stream from Ollama KB model — bypasses provider setting."""
    model = _store.get("ollama_model", "gemma4:12b")
    use_think = "gemma4" in model
    ollama_msgs = _to_ollama_messages(messages)
    if system:
        ollama_msgs = [{"role": "system", "content": system}] + ollama_msgs
    stream = _ollama.chat(model=model, messages=ollama_msgs, stream=True,
        **({"think": False} if use_think else {}),
        options={"num_predict": max_tokens, "num_ctx": num_ctx, "temperature": 0,
                 "keep_alive": -1})
    for chunk in stream:
        token = (chunk.message.content or "")
        if token:
            yield token


def _stream_anthropic(messages: list, max_tokens: int = 400, system: str = ""):
    if not _ANTHROPIC_AVAILABLE:
        raise RuntimeError("anthropic not installed")
    api_key = _store.get("api_key", "").strip()
    if not api_key:
        raise RuntimeError("No Anthropic API key")
    model = _store.get("anthropic_model") or _store["model"]
    client = _anthropic.Anthropic(api_key=api_key)
    kwargs = dict(model=model, max_tokens=max_tokens, temperature=0,
                  messages=_to_anthropic_messages(messages))
    if system:
        kwargs["system"] = system
    with client.messages.stream(**kwargs) as stream:
        for token in stream.text_stream:
            yield token


def _stream_race(messages: list, max_tokens: int = 400, num_ctx: int = 8192, system: str = ""):
    """Fire both Ollama and Anthropic simultaneously; stream tokens from whichever responds first."""
    import queue as _queue
    q = _queue.Queue()
    _SENTINEL = object()

    def _anthropic_worker():
        try:
            for tok in _stream_anthropic(messages, max_tokens, system=system):
                q.put(("a", tok))
        except Exception as e:
            print(f"  [Race/Anthropic] Error: {e}")
            q.put(("a_err", str(e)))
        finally:
            q.put(("a", _SENTINEL))

    def _ollama_worker():
        try:
            for tok in _stream_ollama_kb(messages, max_tokens, num_ctx, system=system):
                q.put(("o", tok))
        except Exception as e:
            print(f"  [Race/Ollama] Error: {e}")
            q.put(("o_err", str(e)))
        finally:
            q.put(("o", _SENTINEL))

    ta = threading.Thread(target=_anthropic_worker, daemon=True)
    to = threading.Thread(target=_ollama_worker, daemon=True)
    ta.start()
    to.start()

    winner = None
    done_count = 0

    while done_count < 2:
        src, tok = q.get()
        if src in ("a_err", "o_err"):
            done_count += 1
            # If the winner errored, we'd get no more tokens from it anyway
            continue
        if tok is _SENTINEL:
            done_count += 1
            continue
        if winner is None:
            winner = src
            label = "Claude" if src == "a" else "Ollama"
            print(f"  [Race] {label} responded first — streaming from {label}")
        if src == winner:
            yield tok
        # discard tokens from the slower model


def _stream(messages: list, max_tokens: int = 400, num_ctx: int = 8192, system: str = ""):
    provider = _store.get("provider")
    if provider == "anthropic":
        yield from _stream_anthropic(messages, max_tokens, system=system)
    elif provider == "hybrid":
        yield from _stream_race(messages, max_tokens, num_ctx, system=system)
    else:
        yield from _stream_ollama(messages, max_tokens, num_ctx, system=system)


def _chat_anthropic(messages: list, max_tokens: int = 1000, system: str = "") -> str:
    if not _ANTHROPIC_AVAILABLE:
        raise RuntimeError("anthropic not installed — run: pip install anthropic")
    api_key = _store.get("api_key", "").strip()
    if not api_key:
        raise RuntimeError("No Anthropic API key — enter it in the Model step")
    model = _store.get("anthropic_model") or _store["model"]
    client = _anthropic.Anthropic(api_key=api_key)
    kwargs = dict(model=model, max_tokens=max_tokens,
                  temperature=0,   # deterministic — same context → same answer
                  messages=_to_anthropic_messages(messages))
    if system:
        kwargs["system"] = system
    resp = client.messages.create(**kwargs)
    text = (resp.content[0].text or "").strip()
    print(f"  [Chat/Anthropic] {len(text)} chars (model={model})")
    return text


def _chat(messages: list, max_tokens: int = 1000, num_ctx: int = 8192, system: str = "") -> str:
    provider = _store.get("provider")
    if provider in ("anthropic", "hybrid"):
        return _chat_anthropic(messages, max_tokens, system=system)
    return _chat_ollama(messages, max_tokens, num_ctx, system=system)


# ── Document parsers ───────────────────────────────────────────────────────────

def _parse_docx(file_obj):
    import zipfile as _zipfile
    from xml.etree import ElementTree as _ET

    NS_B = 'http://schemas.openxmlformats.org/officeDocument/2006/bibliography'
    NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    # ── 1. Build tag → (ref_num, author_str, year, title) from Word citation XML ─
    tag_map   = {}   # tag → (ref_num, "Author (Year)", title)
    order_map = {}   # ref_num → "Author (Year). Title."
    file_obj.seek(0)
    raw_bytes = file_obj.read()
    file_obj.seek(0)

    try:
        with _zipfile.ZipFile(io.BytesIO(raw_bytes)) as z:
            names = z.namelist()
            bib_file = next((n for n in names if n.endswith('item1.xml') and 'customXml' in n), None)
            if bib_file:
                bib_xml  = z.read(bib_file).decode('utf-8')
                bib_root = _ET.fromstring(bib_xml)
                # Detect namespace from actual root tag (robust against version differences)
                root_tag = bib_root.tag  # e.g. '{http://...}Sources'
                bib_ns   = root_tag.split('}')[0].lstrip('{') if '}' in root_tag else NS_B
                print(f"  [DocX] Bibliography file: {bib_file}, namespace: {bib_ns[:40]}...")
                sources_found = bib_root.findall(f'{{{bib_ns}}}Source')
                print(f"  [DocX] Sources found in XML: {len(sources_found)}")
                for source in sources_found:
                    def _bget(el, tag_name, _ns=bib_ns):
                        e = el.find(f'{{{_ns}}}{tag_name}')
                        return e.text.strip() if e is not None and e.text else ''
                    stag   = _bget(source, 'Tag')
                    year   = _bget(source, 'Year')
                    title  = _bget(source, 'Title')
                    order  = _bget(source, 'RefOrder')
                    corp   = source.find(f'.//{{{bib_ns}}}Corporate')
                    persons = source.findall(f'.//{{{bib_ns}}}Person')
                    if corp is not None and corp.text:
                        author = corp.text.strip()
                    elif persons:
                        lasts = []
                        for p in persons:
                            l = p.find(f'{{{bib_ns}}}Last')
                            if l is not None and l.text:
                                lasts.append(l.text.strip())
                        if len(lasts) == 1:   author = lasts[0]
                        elif len(lasts) == 2: author = f"{lasts[0]} y {lasts[1]}"
                        else:                 author = f"{lasts[0]} y colaboradores"
                    else:
                        author = stag
                    ref_num = int(order) if order.isdigit() else 999
                    citation_str  = f"{author} ({year})"
                    tag_map[stag] = (ref_num, citation_str, title)
                    order_map[ref_num] = f"[{ref_num}] {citation_str}. {title}."
                print(f"  [DocX] Loaded {len(tag_map)} bibliography entries from Word XML")
            else:
                print(f"  [DocX] No customXml/item1.xml found. Files: {[n for n in names if 'customXml' in n]}")
    except Exception as exc:
        import traceback; traceback.print_exc()
        print(f"  [DocX] Warning — could not parse bibliography XML: {exc}")

    # ── 2. Rebuild paragraph text with citation tags injected ────────────────────
    def _para_with_citations(para_el):
        """Walk paragraph XML, inserting [Author (Year)] where CITATION fields are."""
        parts     = []
        in_field  = False
        cit_tags  = []
        for elem in para_el.iter():
            local = elem.tag.split('}')[-1]
            if local == 'fldChar':
                ftype = elem.get(f'{{{NS_W}}}fldCharType', '')
                if ftype == 'begin':
                    in_field = True
                    cit_tags = []
                elif ftype == 'end':
                    in_field = False
                    for ct in cit_tags:
                        if ct in tag_map:
                            _, cstr, _ = tag_map[ct]
                            parts.append(f"[{cstr}]")
                        else:
                            parts.append(f"[{ct}]")
                    cit_tags = []
            elif local == 'instrText' and in_field:
                for ct in re.findall(r'CITATION\s+(\S+)', elem.text or ''):
                    cit_tags.append(ct)
            elif local == 't' and not in_field:
                if elem.text:
                    parts.append(elem.text)
        return ''.join(parts).strip()

    # ── 3. Parse document XML for paragraph text (with citations) ────────────────
    sections = []
    try:
        with _zipfile.ZipFile(io.BytesIO(raw_bytes)) as z:
            doc_xml  = z.read('word/document.xml').decode('utf-8')
            doc_root = _ET.fromstring(doc_xml)
            for para_el in doc_root.findall(f'.//{{{NS_W}}}p'):
                t = _para_with_citations(para_el)
                if t:
                    sections.append(t)
    except Exception as exc:
        print(f"  [DocX] Warning — XML paragraph parse failed: {exc}. Falling back to python-docx.")
        doc = DocxDocument(io.BytesIO(raw_bytes))
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                sections.append(t)

    # ── 4. Tables (python-docx for simplicity) ───────────────────────────────────
    doc = DocxDocument(io.BytesIO(raw_bytes))
    for tidx, table in enumerate(doc.tables):
        rows    = []
        headers = []
        for ridx, row in enumerate(table.rows):
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            deduped = []
            for c in cells:
                if not deduped or c != deduped[-1]:
                    deduped.append(c)
            if ridx == 0:
                headers = deduped
            if any(deduped):
                rows.append(" | ".join(deduped))
        if rows:
            header_line = ("COLUMNS: " + " | ".join(headers) + "\n") if headers else ""
            sections.append(f"[TABLE {tidx+1}]\n" + header_line + "\n".join(rows[1:]))

    # ── 5. Append full bibliography as a retrievable section ─────────────────────
    if order_map:
        bib_lines = [order_map[n] for n in sorted(order_map)]
        sections.append("[BIBLIOGRAPHY]\n" + "\n".join(bib_lines))
        print(f"  [DocX] Injected {len(bib_lines)}-entry bibliography section")

    # ── 6. Images ────────────────────────────────────────────────────────────────
    images = []
    try:
        doc2 = DocxDocument(io.BytesIO(raw_bytes))
        for rel in doc2.part.rels.values():
            if "image" in rel.reltype:
                try:
                    part = rel.target_part
                    mt   = part.content_type
                    if mt in ("image/png", "image/jpeg", "image/gif", "image/webp"):
                        images.append({"data": base64.b64encode(part.blob).decode(), "media_type": mt})
                except Exception:
                    pass
    except Exception:
        pass

    # Save tag_map globally so /ask can inject it as a CITATION RESOLVER
    _store["tag_map"] = {tag: cstr for tag, (_, cstr, _title) in tag_map.items()}

    return sections, images


def _parse_pdf(file_obj):
    file_bytes = file_obj.read()
    sections   = []
    images     = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                for line in text.split("\n"):
                    line = line.strip()
                    if line:
                        sections.append(line)
            for table in page.extract_tables():
                rows    = []
                headers = []
                for ridx, row in enumerate(table):
                    cells   = [c.strip() if c else "" for c in row]
                    deduped = []
                    for c in cells:
                        if not deduped or c != deduped[-1]:
                            deduped.append(c)
                    if ridx == 0:
                        headers = deduped
                    if any(deduped):
                        rows.append(" | ".join(deduped))
                if rows:
                    header_line = ("COLUMNS: " + " | ".join(headers) + "\n") if headers else ""
                    sections.append("[TABLE]\n" + header_line + "\n".join(rows[1:]))
    if _FITZ_AVAILABLE:
        try:
            doc = _fitz.open(stream=file_bytes, filetype="pdf")
            seen = set()
            for page in doc:
                for img_info in page.get_images(full=True):
                    xref = img_info[0]
                    if xref in seen:
                        continue
                    seen.add(xref)
                    try:
                        base = doc.extract_image(xref)
                        ext  = base["ext"]
                        mt   = f"image/{'jpeg' if ext == 'jpg' else ext}"
                        if mt in ("image/png", "image/jpeg", "image/gif", "image/webp"):
                            images.append({"data": base64.b64encode(base["image"]).decode(), "media_type": mt})
                    except Exception:
                        pass
            doc.close()
        except Exception as e:
            print(f"  [PDF] Image extraction warning: {e}")
    return sections, images


# ── Keyword index ──────────────────────────────────────────────────────────────

def _normalize_spanish(words: set) -> set:
    expanded = set(words)
    for w in list(words):
        if len(w) > 5 and w.endswith("es"):
            expanded.add(w[:-2])
        if len(w) > 4 and w.endswith("s"):
            expanded.add(w[:-1])
    return expanded


def _build_index(paragraphs: list) -> dict:
    index = {}
    for i, para in enumerate(paragraphs):
        words = _normalize_spanish(set(re.findall(r"\w{3,}", para.lower())))
        for word in words:
            index.setdefault(word, set()).add(i)
    return index

_FILLER_PREFIXES = re.compile(
    r"^(?:"
    r"what would you say about|what do you know about|what can you tell me about"
    r"|can you (?:please )?(?:explain|describe|tell me about)"
    r"|could you (?:please )?(?:explain|describe|tell me about)"
    r"|i(?:'d| would) like to know(?: about)?"
    r"|i (?:want|need) to know(?: about)?"
    r"|tell me(?: all)? about|tell me"
    r"|explain(?: to)?(?: me)?"
    r"|describe(?: to me)?"
    r"|talk(?: to me)? about"
    r"|give me (?:info|information|details)?(?:\s*(?:about|on|regarding))?"
    r"|show me"
    r"|what(?:'s| is| are)?(?: the)?"
    r"|how(?: about| does| do| is| are)?"
    r"|qu[eé] (?:es|son|sabes de|me puedes decir(?: sobre)?)"
    r"|expl[ií]came(?: sobre)?"
    r"|h[aá]blame(?: de| sobre)?"
    r"|cu[eé]ntame(?: (?:sobre|acerca de))?"
    r"|dime(?: algo)?(?: (?:sobre|acerca de))?"
    r"|dame (?:informaci[oó]n|info|detalles)?(?:\s*(?:sobre|de|acerca de))?"
    r"|describe"
    r"|(?:me|nos) (?:pueden?|podr[ií]an?|gustar[ií]a) (?:por favor )?(?:explicar|decir|contar|describir)(?: qu[eé])?"
    r"|(?:me|nos) (?:pueden?|podr[ií]an?) (?:por favor )?"
    r")\s+",
    re.IGNORECASE,
)
_FILLER_SUFFIXES  = re.compile(r"\s*(?:please|por favor)\s*\??$|\?$", re.IGNORECASE)
_LEADING_ARTICLES = re.compile(r"^(?:the|a|an|el|la|los|las|un|una|about|sobre|de)\s+", re.IGNORECASE)


def _extract_topic(question: str) -> str:
    """Strip conversational filler to expose core topic keywords. No LLM call."""
    text = question.strip()
    for _ in range(2):
        cleaned = _FILLER_PREFIXES.sub("", text).strip()
        if cleaned == text:
            break
        text = cleaned
    text = _FILLER_SUFFIXES.sub("", text).strip()
    text = _LEADING_ARTICLES.sub("", text).strip()
    result = text if text else question
    if result != question:
        print(f"  [Topic] '{question}' → '{result}'")
    return result

# Hardcoded semantic clusters for academic/cybersecurity terms that
# the auto-generated concept map frequently misses.
_QUERY_CLUSTERS = {
    # methodology cluster — catches PDCA chapter content
    "metodolog": ["pdca", "planificar", "hacer", "verificar", "actuar",
                  "ciclo", "mejora", "continua", "deming", "plan", "check",
                  "do", "act", "fase", "etapa", "proceso", "enfoque",
                  "investigacion", "experimental", "validacion", "diseño"],
    "pdca":      ["metodolog", "planificar", "hacer", "verificar", "actuar",
                  "ciclo", "mejora", "continua", "deming"],
    # results / comparison cluster
    "resultado": ["hallazgo", "obtuv", "logr", "alcanz", "medicion",
                  "rendimiento", "latencia", "disponibilidad", "throughput"],
    "comparar":  ["estudio", "autor", "literatura", "previo", "similar",
                  "diferencia", "coincid", "contraste", "versus"],
    # architecture cluster
    "arquitect": ["sdwan", "starlink", "vpn", "firewall", "nodo", "enlace",
                  "topolog", "red", "infraestructur", "diseno"],
    "sdwan":     ["arquitect", "wan", "enlace", "nodo", "topolog", "red"],
    # risk / security cluster
    "riesgo":    ["amenaza", "vulnerabilid", "control", "mitigacion",
                  "iso", "nist", "ciberseguridad", "gestion"],
    "implementa": ["despleg", "configurar", "instalar", "integrar",
                   "desarroll", "construir", "ejecutar", "propuesta"],
}


def _expand_query(question: str) -> list:
    cm       = _store.get("concept_map", {})
    keywords = _normalize_spanish(set(re.findall(r"\w{3,}", question.lower())))
    expanded = set(keywords)

    # 1. Concept-map expansion (auto-generated per document)
    for alt, canonical in cm.get("aliases", {}).items():
        if isinstance(canonical, list):
            canonical = " ".join(str(x) for x in canonical)
        if set(re.findall(r"\w{3,}", alt.lower())) & keywords:
            expanded.update(re.findall(r"\w{3,}", canonical.lower()))
    for section in cm.get("sections", []):
        if isinstance(section, list):
            section = " ".join(str(x) for x in section)
        sec_words = set(re.findall(r"\w{3,}", section.lower()))
        if sec_words & keywords:
            expanded.update(sec_words)

    # 2. Hardcoded cluster expansion — catches terms the concept map misses
    for trigger, extras in _QUERY_CLUSTERS.items():
        if any(kw.startswith(trigger) or trigger.startswith(kw) for kw in keywords):
            expanded.update(extras)

    if expanded != keywords:
        print(f"  [Expand] added: {sorted(expanded - keywords)}")
    return list(expanded)


def _retrieve(keywords: list, top_n: int = 60) -> tuple:
    paragraphs = _store["paragraphs"]
    raw_count  = _store["raw_count"]
    index      = _store["index"]
    all_table_indices = [i for i, p in enumerate(paragraphs[:raw_count])
                         if p.startswith("[TABLE")]

    scores: dict[int, int] = {}
    for kw in keywords:
        for idx in index.get(kw, set()):
            scores[idx] = scores.get(idx, 0) + 1
    if not scores:
        table_text = "\n\n".join(paragraphs[i] for i in all_table_indices)
        return "\n".join(paragraphs[:20]) + ("\n\n" + table_text if table_text else ""), ""
    top     = sorted(scores, key=scores.__getitem__, reverse=True)[:top_n]
    top     = sorted(top)
    header  = list(range(min(5, raw_count) if raw_count else min(5, len(paragraphs))))
    indices = sorted(set(header + top))

    # TABLE EXPANSION
    expanded = set(indices)
    table_found = False
    for i in list(indices):
        if i >= raw_count:
            continue
        start = i
        while start > 0 and not paragraphs[start].startswith("[TABLE"):
            start -= 1
        if paragraphs[start].startswith("[TABLE"):
            table_found = True
            end = start + 1
            while end < raw_count and not paragraphs[end].startswith("[TABLE"):
                end += 1
            expanded.update(range(start, end))

    indices  = sorted(expanded)
    # Only merge all_table_indices when the query actually matched table content;
    # merging unconditionally caused has_tables_in_context to always be True,
    # which suppressed the KB for every question regardless of relevance.
    matched_table_indices = [i for i in all_table_indices if i in set(indices)]
    indices_with_tables   = sorted(set(indices) | set(matched_table_indices))

    # Separate table rows from regular paragraphs so tables go first in the
    # context window — they were getting truncated when large docs pushed them
    # past the 10000-char cutoff.
    table_parts  = [paragraphs[i] for i in indices_with_tables if i < raw_count and paragraphs[i].startswith("[TABLE")]
    other_parts  = [paragraphs[i] for i in indices_with_tables if i < raw_count and not paragraphs[i].startswith("[TABLE")]
    kb_parts     = [paragraphs[i] for i in sorted(set(indices)) if i >= raw_count]

    has_tables_in_context = bool(table_parts)

    # Find the [BIBLIOGRAPHY] section injected by _parse_docx (Word citation XML)
    bib_idx = next(
        (i for i, p in enumerate(paragraphs[:raw_count]) if p.startswith("[BIBLIOGRAPHY]")),
        None,
    )
    if bib_idx is None:
        bib_idx = next(
            (i for i, p in enumerate(paragraphs[:raw_count]) if re.search(r"^\[\d+\]", p.strip())),
            None,
        )
    ref_block = (
        "\n\n---REFERENCES---\n" + paragraphs[bib_idx][:5000]
    ) if bib_idx is not None else ""

    # Tables first so they're never cut off, then remaining paragraphs up to budget.
    # 5000-char cap (≈1250 tokens) keeps prompt encoding fast on a local 12B model.
    RAW_BUDGET = 5000
    table_text = "\n\n".join(table_parts)
    other_text = "\n".join(other_parts)
    remaining  = max(0, RAW_BUDGET - len(table_text) - 2)
    raw_text   = (table_text + "\n\n" + other_text[:remaining]).strip() if table_text else other_text[:RAW_BUDGET]

    if table_found or has_tables_in_context:
        return raw_text + ref_block[:2000], ""

    return raw_text + ref_block[:2000], "\n".join(kb_parts)[:2000]

# ── Background processing ──────────────────────────────────────────────────────

_KB_PROMPT_TEMPLATE = (
    "Build a complete reference knowledge base from this document{part_note}.\n\n"
    "MANDATORY REQUIREMENTS:\n"
    "1. List every section/chapter title with its exact name.\n"
    "2. For each section: include ALL content — every fact, number, "
    "name, list item, and table row. Omit nothing.\n"
    "3. TABLES — ABSOLUTE RULE: reproduce the table as a numbered list.\n"
    "   Each row becomes one numbered item using the EXACT cell text.\n"
    "   FORBIDDEN: rewording, expanding, merging, or explaining any row.\n"
    "   Example — if table has: 'Dependencia de enlace único'\n"
    "   You write: '1. Dependencia de enlace único'\n"
    "   NOT: '1. Riesgo de depender de un solo enlace de comunicación'\n"
    "   Count rows before writing. If 10 rows → exactly 10 numbered items.\n"
    "4. Images/charts: describe their content in full detail.\n"
    "5. Preserve all proper nouns, numbers, and technical terms exactly as written.\n"
    "6. If there is a Q&A section or defense preparation questions with answers, "
    "   reproduce them VERBATIM — do not summarize or paraphrase.\n\n"
    "DOCUMENT:\n{chunk}"
)


def _process_in_background(text: str, images: list):
    """
    Builds a knowledge base from the document using the selected model.
    In Hybrid or Anthropic+Ollama modes, also runs Ollama KB in parallel for richer retrieval.
      1. A verbatim knowledge base (all sections, all table rows, Q&A, image descriptions)
      2. A concept map (section names + spoken-language aliases)
    """
    _store["processing"]  = True
    _store["proc_error"]  = ""
    _store["summary"]     = ""
    _store["concept_map"] = {}

    provider = _store.get("provider", "ollama")
    # Gemma 4 12B: 256K token window. Send up to ~140K chars per chunk (~35K tokens).
    MAX_CHARS = 140_000
    chunks    = [text[i:i + MAX_CHARS] for i in range(0, len(text), MAX_CHARS)]
    summaries = []

    # ── Parallel Ollama KB: run when using Anthropic or Hybrid so Gemma also
    #    indexes the document (Gemma is better at literal Q&A extraction).
    ollama_extra_summaries: list = []
    ollama_kb_thread = None
    def _ollama_is_running():
        try:
            _ollama.list()
            return True
        except Exception:
            return False

    run_parallel_ollama = (
        _store.get("ollama_model")
        and provider in ("anthropic", "hybrid")
        and _ollama_is_running()
    )
    if run_parallel_ollama:
        def _ollama_kb_worker():
            for cidx, chunk in enumerate(chunks):
                part_note = f" (part {cidx+1} of {len(chunks)})" if len(chunks) > 1 else ""
                content = [{"type": "text", "text": _KB_PROMPT_TEMPLATE.format(
                    part_note=part_note, chunk=chunk)}]
                try:
                    result = _chat_ollama_kb(
                        [{"role": "user", "content": content}],
                        max_tokens=10000,
                        num_ctx=40000,
                    )
                    ollama_extra_summaries.append(result)
                    print(f"  [Hybrid/OllamaKB] Chunk {cidx+1} done: {len(result):,} chars")
                except Exception as e:
                    print(f"  [Hybrid/OllamaKB] Chunk {cidx+1} error: {e}")
        ollama_kb_thread = threading.Thread(target=_ollama_kb_worker, daemon=True)
        ollama_kb_thread.start()
        print(f"  [Hybrid] Ollama KB building started in parallel (model={_store.get('ollama_model')})")

    try:
        for cidx, chunk in enumerate(chunks):
            part_note = f" (part {cidx+1} of {len(chunks)})" if len(chunks) > 1 else ""
            content = [{"type": "text", "text": _KB_PROMPT_TEMPLATE.format(
                part_note=part_note, chunk=chunk)}]
            # Attach images to first chunk (Ollama vision: image_url format)
            if cidx == 0 and images:
                for img in images[:15]:
                    content.append({
                        "type": "image_url",
                        "image_url": {"url": f"data:{img['media_type']};base64,{img['data']}"},
                    })
                content.append({
                    "type": "text",
                    "text": "The images above are from the document. Describe their content and include it under the relevant section.",
                })

            print(f"  [Processing] Chunk {cidx+1}/{len(chunks)}: {len(chunk):,} chars → generating KB…")
            # Anthropic models cap at 8192 output tokens; Ollama supports 10000+
            # Both anthropic and hybrid run the Anthropic KB pass here.
            # Ollama KB runs in parallel (started above) and gets merged after.
            proc_max_tokens = 8000 if provider in ("anthropic", "hybrid") else 10000
            result = _chat([{"role": "user", "content": content}], max_tokens=proc_max_tokens, num_ctx=40000)
            summaries.append(result)
            print(f"  [Processing] Chunk {cidx+1} done: {len(result):,} chars")

        # Wait for parallel Ollama KB
        if ollama_kb_thread is not None:
            print("  [Hybrid] Waiting for Ollama KB to finish…")
            ollama_kb_thread.join(timeout=600)
            if ollama_extra_summaries:
                combined = "\n\n---\n\n".join(ollama_extra_summaries)
                extra_paras = [p.strip() for p in combined.split("\n") if len(p.strip()) > 20]
                print(f"  [Hybrid] Ollama KB ready: {len(extra_paras)} paragraphs")
                summaries = summaries + ollama_extra_summaries  # merge both KBs
            else:
                print("  [Hybrid] Ollama KB produced no output — using primary KB only")

        if not summaries:
            raise RuntimeError("No KB content generated — check model availability")

        _store["summary"] = "\n\n---\n\n".join(summaries)
        print(f"  [Processing] Knowledge base: {len(_store['summary']):,} chars total")

        # Merge KB paragraphs into the searchable index alongside raw paragraphs.
        # raw_count marks the boundary — raw paragraphs take priority for exact table content.
        kb_paras = [p.strip() for p in _store["summary"].split("\n") if len(p.strip()) > 20]
        _store["paragraphs"] = _store["paragraphs"][:_store["raw_count"]] + kb_paras
        _store["index"]      = _build_index(_store["paragraphs"])
        print(f"  [Processing] Index: {_store['raw_count']} raw + {len(kb_paras)} KB paragraphs")

        # Build concept map from KB
        print("  [Processing] Building concept map…")
        try:
            cm_result = _chat(
                messages=[{"role": "user", "content": (
                    "Extract a vocabulary index from this document reference. "
                    "Return strict JSON only — no other text:\n"
                    '{"sections": ["exact section title", ...], '
                    '"aliases": {"alternative term": "document term", ...}}\n\n'
                    "For aliases: include synonyms someone might say out loud.\n"
                    "Examples: doc has 'Risks' → aliases: barriers, problems, issues\n"
                    "          doc has 'Planning Phase' → aliases: implementation, execution\n"
                    "IMPORTANT: document is in Spanish but users may ask in English. "
                    "Add English translations for key terms. "
                    "Example: 'riesgos' → risks, risk\n"
                    "         'planificación' → planning, planning phase\n"
                    "Include accented and unaccented variants.\n\n"
                    f"REFERENCE:\n{_store['summary'][:6000]}"
                )}],
                max_tokens=1500,
                num_ctx=10000,   # 6K reference + 1.5K output + buffer
            )
            match = re.search(r"\{[\s\S]*\}", cm_result)
            if match:
                cm = json.loads(match.group())
                _store["concept_map"] = cm
                print(f"  [Processing] Concept map: {len(cm.get('sections',[]))} sections, "
                      f"{len(cm.get('aliases',{}))} aliases")
            else:
                print("  [Processing] Concept map: no JSON found, continuing without")
        except Exception as e:
            print(f"  [Processing] Concept map warning: {e}")

        _store["processing"] = False
        print("  [Processing] ✓ Done — Start Listening is now available")

    except Exception as exc:
        _store["proc_error"] = str(exc)
        _store["processing"] = False
        print(f"  [Processing] ERROR: {exc}")


# ── Transcription ──────────────────────────────────────────────────────────────

def _vocab_hint() -> str:
    """Build a short vocabulary prompt for Whisper from the concept map."""
    cm    = _store.get("concept_map", {})
    parts = []
    # Section names first — most useful for domain recognition
    if cm.get("sections"):
        for s in cm["sections"][:8]:
            if isinstance(s, str):
                parts.append(s)
            elif isinstance(s, list):
                parts.extend(s)

    # Then key document terms (aliases values can be str or list)
    for v in list(cm.get("aliases", {}).values())[:12]:
        if isinstance(v, str):
            parts.append(v)
        elif isinstance(v, list):
            parts.extend(str(x) for x in v)

    return ", ".join(parts)


def _wav_to_float32(audio_bytes: bytes) -> np.ndarray:
    """Decode WAV bytes → float32 numpy array in range [-1, 1]."""
    with wave.open(io.BytesIO(audio_bytes)) as wf:
        raw    = wf.readframes(wf.getnframes())
        n_ch   = wf.getnchannels()
        sw     = wf.getsampwidth()
    if sw == 2:
        arr = np.frombuffer(raw, dtype=np.int16).astype(np.float32) / 32768.0
    elif sw == 4:
        arr = np.frombuffer(raw, dtype=np.int32).astype(np.float32) / 2147483648.0
    else:
        arr = np.frombuffer(raw, dtype=np.float32)
    if n_ch > 1:
        arr = arr.reshape(-1, n_ch).mean(axis=1)
    return arr


def _transcribe_whisper(audio_bytes: bytes, lang: str, model_name: str = "medium") -> str:
    """
    Transcribe using faster-whisper locally.
    - No network roundtrip: ~0.3s per 3s clip on CPU (int8)
    - initial_prompt seeds Whisper with document vocabulary so technical
      terms like VPN, SECOP, Peplink are recognised correctly.
    """
    lang_code = lang.split("-")[0]   # "es-ES" → "es"
    vocab     = _vocab_hint()
    arr       = _wav_to_float32(audio_bytes)
    model     = _get_whisper(model_name)
    print(f"  [Whisper] audio: {len(arr)} samples, duration={len(arr)/16000:.1f}s, rms={float(np.sqrt(np.mean(arr**2))):.5f}, max={float(np.max(np.abs(arr))):.4f}")

    segments, info = model.transcribe(
        arr,
        language=lang_code,
        task="transcribe",
        initial_prompt=vocab if vocab else None,
        beam_size=1,
        condition_on_previous_text=False,   # skip recurrent context — faster
        vad_filter=False,
    )
    text = " ".join(seg.text.strip() for seg in segments).strip()
    print(f"  [Whisper] {info.language} ({info.language_probability:.0%}): {text!r}")
    return text


def _transcribe_google(audio_bytes: bytes, lang: str) -> str:
    if not _SR_AVAILABLE:
        raise RuntimeError("SpeechRecognition not installed — run: pip install SpeechRecognition")
    with sr.AudioFile(io.BytesIO(audio_bytes)) as source:
        audio = _sr.record(source)
    return _sr.recognize_google(audio, language=lang)


# ── Routes ─────────────────────────────────────────────────────────────────────

@app.route("/")
def index():
    return Response(INDEX_HTML, mimetype="text/html")

@app.route("/favicon.ico")
def favicon():
    return "", 204


@app.route("/load-document", methods=["POST"])
def load_document():
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400

    f            = request.files["file"]
    filename     = f.filename or ""
    ext          = Path(filename).suffix.lower()
    model        = request.form.get("model", DEFAULT_MODEL).strip()
    provider     = request.form.get("provider", "ollama").strip()
    api_key      = request.form.get("api_key", "").strip()
    ollama_model = request.form.get("ollama_model", "").strip()
    anthropic_model = request.form.get("anthropic_model", "").strip()
    if model:           _store["model"]           = model
    if provider:        _store["provider"]        = provider
    if api_key:         _store["api_key"]         = api_key
    if ollama_model:    _store["ollama_model"]    = ollama_model
    if anthropic_model: _store["anthropic_model"] = anthropic_model
    # For hybrid: model field is the Anthropic answer model
    if provider == "hybrid" and anthropic_model:
        _store["model"] = anthropic_model

    try:
        if ext == ".docx":
            sections, images = _parse_docx(f)
        elif ext == ".pdf":
            sections, images = _parse_pdf(f)
        else:
            return jsonify({"error": "Unsupported file. Use .docx or .pdf"}), 400
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500

    text = "\n".join(sections)
    _store.update({
        "doc_text":   text,
        "doc_name":   filename,
        "paragraphs": sections,
        "raw_count":  len(sections),
        "index":      _build_index(sections),
        "images":     images,
        "summary":    "",
        "concept_map": {},
        "proc_error": "",
    })

    # Mark processing=True BEFORE starting the thread to close the race window
    # where the UI could poll /process-status before the thread sets it.
    _store["processing"] = True

    t = threading.Thread(target=_process_in_background, args=(text, images), daemon=True)
    t.start()
    print(f"  [Load] {filename}  {len(text):,} chars, {len(images)} images, model={_store['model']}")

    return jsonify({
        "name":       filename,
        "chars":      len(text),
        "sections":   len(sections),
        "images":     len(images),
        "processing": True,
    })


@app.route("/process-status")
def process_status():
    # "done" only fires when BOTH the KB and concept map are complete.
    # Checking both conditions prevents the race where the UI unlocks
    # while the concept map job is still holding the Ollama GPU.
    done = bool(_store["summary"]) and not _store["processing"]
    return jsonify({
        "processing": _store["processing"],
        "done":       done,
        "error":      _store["proc_error"],
    })


_whisper_cache: dict = {}   # model_name → WhisperModel instance

def _get_whisper(model_name: str):
    """Return a cached WhisperModel, loading on first use."""
    if model_name not in _whisper_cache:
        print(f"  [Whisper] Loading {model_name} on CPU (int8) — first-use delay…")
        _whisper_cache[model_name] = _WhisperModel(model_name, device="cpu", compute_type="int8")
        print(f"  [Whisper] {model_name} ready")
    return _whisper_cache[model_name]


@app.route("/transcribe", methods=["POST"])
def transcribe():
    lang        = request.args.get("lang", "en-US")
    mode        = request.args.get("mode", "whisper")   # "whisper" | "google" | "auto"
    wmodel      = request.args.get("wmodel", "medium")  # whisper model size
    audio_bytes = request.get_data()

    if not audio_bytes or len(audio_bytes) < 1000:
        return jsonify({"text": ""})

    try:
        if mode == "google":
            text = _transcribe_google(audio_bytes, lang)

        elif mode in ("whisper", "auto"):
            if _WHISPER_AVAILABLE:
                text = _transcribe_whisper(audio_bytes, lang, model_name=wmodel)
            else:
                text = _transcribe_google(audio_bytes, lang)

        else:
            return jsonify({"error": f"Unknown transcription mode: {mode}"}), 400

    except Exception as e:
        if "UnknownValueError" in type(e).__name__:
            return jsonify({"text": ""})
        return jsonify({"error": str(e)}), 500

    return jsonify({"text": text or ""})



@app.route("/ask", methods=["POST"])
def ask():
    data     = request.get_json(force=True)
    question = (data.get("question") or "").strip()
    language = (data.get("language") or "Spanish")
    model    = (data.get("model") or "").strip()
    provider = (data.get("provider") or "").strip()
    api_key  = (data.get("api_key") or "").strip()
    if model:    _store["model"]    = model
    if provider: _store["provider"] = provider
    if api_key:  _store["api_key"]  = api_key

    if not question:
        return jsonify({"error": "No question provided"}), 400
    if not _store["paragraphs"]:
        return jsonify({"error": "No document loaded"}), 400
    if _store["processing"]:
        return jsonify({"error": "Document still processing, please wait"}), 503

    topic    = _extract_topic(question)
    keywords = _expand_query(topic)
    # Wider retrieval for deep questions (methodology, studies, comparisons)
    _deep_triggers = (
        "metodolog", "pdca", "estudio", "literatura", "comparar", "resultado",
        "hallazgo", "revision", "investigaci", "bibliograf", "marco", "teoric"
    )
    _retrieval_top_n = 100 if any(t in question.lower() for t in _deep_triggers) else 60
    raw_text, kb_text = _retrieve(keywords, top_n=_retrieval_top_n)

    context = "---SOURCE DOCUMENT (authoritative)---\n" + raw_text
    if kb_text:
        context += "\n\n---PROCESSED REFERENCE (additional context)---\n" + kb_text

    cm           = _store.get("concept_map", {})
    section_list = " | ".join(cm.get("sections", [])[:20])

    system_prompt = (
        "You are one of three co-authors defending your master's thesis in cybersecurity. "
        "Speak as a human researcher — NEVER as an AI. Use WE/NOSOTROS always "
        "('nosotros implementamos', 'nuestra propuesta', 'encontramos'). "
        "Start immediately with the answer — no preamble.\n\n"
        "FORMAT: Flowing prose only. No bullets, lists, dashes, or markdown. "
        "Simple fact → 1-3 sentences. Complex → 2-3 short paragraphs max. "
        "Exact technical terms: SD-WAN, Starlink, AES-256, IPSec, IKEv2, ISO 27001, SpeedFusion, Peplink. "
        "Silently fix speech-to-text errors. "
        "FORBIDDEN filler: 'El documento menciona', 'Según el texto', 'according to', 'I hope this helps'.\n\n"
        "TECHNICAL PRECISION: 'qué tipo/protocolo/algoritmo/cuánto' → give the SPECIFIC named value, never describe what it does.\n\n"
        "CITATIONS: ONLY cite entries from ---CITATION RESOLVER---. "
        "Never invent authors/years. Never write raw tag codes (And20, GRA22, etc.) — resolve them via the table. "
        "Say 'y colaboradores' not 'et al.'. If unsure, omit the citation entirely — hallucinated citations are worse than none.\n\n"
        "STUDIES: When asked about literature/comparisons, name EVERY relevant study: "
        "'García y colaboradores (2021) demostraron X; López (2020) encontró Y.' Never vague summaries.\n\n"
        "TABLES: Read [TABLE N] row values and speak them as prose. "
        "WRONG PREMISES: Correct them and answer with the right info. "
        "Missing info: one direct sentence with 'nosotros'.\n\n"
        "CLARIFICATION EXCHANGES: The question may contain a short back-and-forth conversation "
        "(jury question → your clarifying question → jury's refined answer). "
        "In that case, identify the FINAL refined scope and answer only that. "
        "Example: 'riesgos del proyecto / los de planificacion / la de planificacion' → answer risks from the planning phase only."
    )

    # Build CITATION RESOLVER and append to system prompt so the model treats it as a hard constraint
    tag_map_store = _store.get("tag_map", {})
    if tag_map_store:
        resolver_lines = "\n".join(f"  {tag} → {cstr}" for tag, cstr in sorted(tag_map_store.items()))
        system_prompt += (
            f"\n\n---CITATION RESOLVER (authoritative — only these citations exist)---\n"
            f"{resolver_lines}\n"
            f"RULE: Only cite entries from this table. Any author or year NOT in this table is HALLUCINATED — omit it."
        )

    user_msg = (
        f"Available sections: {section_list}\n\n"
        f"{context}\n\n"
        f"Question: {question}\n"
        f"Respond in {language}."
    )

    # Size the context window to just what's needed — smaller num_ctx = faster TTFT
    total_chars = len(system_prompt) + len(user_msg)
    estimated_tokens = total_chars // 3  # conservative (Spanish avg ~3 chars/token)
    # Round up to next multiple of 1024, minimum 4096, maximum 12288
    raw_ctx = max(4096, min(12288, ((estimated_tokens + 1500) // 1024 + 1) * 1024))
    print(f"  [Ask] ctx={len(raw_text)+len(kb_text)} chars, tokens≈{estimated_tokens}, num_ctx={raw_ctx}, q={question!r}")

    def generate():
        try:
            full = []
            # Study/comparison/methodology questions need more tokens
            study_keywords = (
                "estudio", "autor", "marco teórico", "investigaci", "trabajo",
                "referencia", "bibliograf", "literatura", "revisión", "revision",
                "comparan", "compara", "comparar", "comparación", "comparacion",
                "antecedente", "previo", "resultado", "hallazgo", "aporte",
                "diferencia", "similar", "coincid", "contrast"
            )
            methodology_keywords = (
                "metodolog", "pdca", "planificar", "ciclo", "mejora",
                "enfoque", "diseño", "proceso", "etapa", "fase", "validaci"
            )
            is_study_q = any(w in question.lower() for w in study_keywords)
            is_method_q = any(w in question.lower() for w in methodology_keywords)
            answer_max_tokens = 1200 if (is_study_q or is_method_q) else 450
            for token in _stream([{"role": "user", "content": user_msg}],
                                  max_tokens=answer_max_tokens, num_ctx=raw_ctx, system=system_prompt):
                full.append(token)
                yield f"data: {json.dumps(token)}\n\n"
            print(f"  [Ask/stream] {sum(len(t) for t in full)} chars")
            yield "data: [DONE]\n\n"
        except Exception as exc:
            import traceback; traceback.print_exc()
            err_payload = json.dumps({"error": str(exc)})
            yield f"data: {err_payload}\n\n"
            yield "data: [DONE]\n\n"

    return Response(generate(), mimetype="text/event-stream",
                    headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


# ── Embedded UI ────────────────────────────────────────────────────────────────

INDEX_HTML = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Meeting Assistant v3</title>
<style>
  *, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }
  :root {
    --bg:#16161e; --panel:#1e1e2a; --card:#24243a; --border:#2e2e4a;
    --text:#e0e0f0; --dim:#7070a0;
    --green:#3ddc84; --green-dim:#1a3d2b; --green-dark:#2a6a4a;
    --blue:#5b9cf6; --blue-dim:#1a2a4a;
    --yellow:#f0c060; --red:#f06060; --step-bg:#2a2a4a;
    --purple:#b57bee; --purple-dim:#2a1a4a;
  }
  body { background:var(--bg); color:var(--text); font-family:'Segoe UI',system-ui,sans-serif;
    font-size:14px; min-height:100vh; display:flex; flex-direction:column;
    align-items:center; padding:24px 16px 40px; }
  .container { width:100%; max-width:720px; display:flex; flex-direction:column; }
  .title-bar { display:flex; align-items:baseline; gap:10px; margin-bottom:12px; }
  .title-bar h1 { font-size:20px; font-weight:700; }
  .v2tag { font-size:12px; background:var(--purple-dim); color:var(--purple); border-radius:4px; padding:2px 7px; font-weight:700; }
  .title-bar .sub { font-size:13px; color:var(--dim); }
  hr { border:none; border-top:1px solid var(--border); margin-bottom:16px; }
  .section-label { font-size:10px; font-weight:700; letter-spacing:.08em; color:var(--dim); margin-bottom:6px; }
  .step { display:flex; align-items:center; background:var(--card); border:1px solid var(--border);
    border-radius:6px; margin-bottom:4px; min-height:52px; overflow:hidden; }
  .step-badge { background:var(--step-bg); color:var(--blue); font-size:15px; font-weight:700;
    width:44px; min-width:44px; align-self:stretch; display:flex; align-items:center; justify-content:center; }
  .step-label { font-weight:700; font-size:13px; min-width:90px; padding-left:10px; }
  .step-content { flex:1; display:flex; align-items:center; gap:8px; padding:8px; flex-wrap:wrap; }
  .step-indicator { padding:0 12px; font-size:14px; color:var(--dim); min-width:30px; text-align:right; }
  .step-indicator.ok { color:var(--green); font-size:16px; }
  input[type="text"], select {
    background:var(--panel); border:none; border-radius:4px; color:var(--text);
    font-family:inherit; font-size:13px; padding:7px 10px; outline:none; transition:box-shadow .15s; }
  input:focus, select:focus { box-shadow:0 0 0 2px var(--blue); }
  input::placeholder { color:var(--dim); }
  select { cursor:pointer; } select option { background:var(--panel); }
  .btn { border:none; border-radius:4px; cursor:pointer; font-family:inherit;
    font-size:13px; font-weight:700; padding:7px 14px; transition:opacity .15s,filter .15s; }
  .btn:hover:not(:disabled) { filter:brightness(1.1); }
  .btn:disabled { opacity:.45; cursor:default; }
  .btn-blue  { background:var(--blue-dim); color:var(--blue); }
  .btn-dim   { background:var(--card); color:var(--dim); }
  .btn-green { background:var(--green); color:#0a1a10; font-size:15px; padding:13px; }
  .btn-red   { background:var(--red); color:#fff; font-size:15px; padding:13px; }
  #start-btn { width:100%; border-radius:6px; margin-top:8px; }
  .doc-name { flex:1; font-size:13px; color:var(--dim); white-space:nowrap; overflow:hidden; text-overflow:ellipsis; }
  .doc-name.loaded { color:var(--text); }
  .status-bar { display:flex; align-items:center; gap:8px; margin:12px 0 8px; }
  .status-dot { font-size:16px; color:var(--dim); }
  .status-text { font-size:13px; color:var(--dim); }
  .tag { font-size:11px; padding:2px 7px; border-radius:3px; font-weight:700; background:var(--purple-dim); color:var(--purple); }
  .panel-header { display:flex; align-items:center; justify-content:space-between; margin-bottom:4px; }
  #transcript-box { background:var(--card); border-radius:6px; color:var(--dim);
    font-size:13px; line-height:1.5; min-height:70px; max-height:90px;
    overflow-y:auto; padding:10px 12px; margin-bottom:14px; word-break:break-word; }
  .answer-header { display:flex; align-items:baseline; gap:12px; margin-bottom:6px; }
  #question-label { font-size:13px; font-style:italic; color:var(--blue); flex:1; }
  #answer-card { background:var(--green-dim); border:1px solid var(--green); border-radius:6px;
    flex:1; min-height:120px; padding:16px; overflow-y:auto; }
  #answer-text { font-size:16px; line-height:1.55; color:var(--green-dark); white-space:pre-wrap; word-break:break-word; }
  #answer-text.has-answer { color:var(--green); }
  ::-webkit-scrollbar { width:6px; } ::-webkit-scrollbar-track { background:transparent; }
  ::-webkit-scrollbar-thumb { background:var(--border); border-radius:3px; }
  #file-input { display:none; }
  .answer-section { display:flex; flex-direction:column; flex:1; }
  .vol-bar-wrap { display:flex; align-items:center; gap:6px; }
  .vol-bar-track { background:var(--border); border-radius:3px; height:6px; width:80px; overflow:hidden; }
  .vol-bar-fill  { background:var(--green); height:100%; width:0%; transition:width .1s; border-radius:3px; }
  .tgl-group { display:flex; gap:4px; }
  .tgl { background:var(--step-bg); color:var(--dim); border:1px solid var(--border);
    border-radius:4px; padding:4px 10px; cursor:pointer; font-size:12px; font-weight:600;
    font-family:inherit; transition:all .15s; }
  .tgl.active { background:var(--purple-dim); color:var(--purple); border-color:var(--purple); }
</style>
</head>
<body>
<div class="container">
  <div class="title-bar">
    <h1>Meeting Assistant</h1>
    <span class="v2tag">v3</span>
    <span class="sub">Gemma 4 / Sonnet · streaming · CPU Whisper</span>
  </div>
  <hr>
  <div class="section-label">SETUP</div>

  <div class="step">
    <div class="step-badge">1</div>
    <div class="step-label">Model</div>
    <div class="step-content" style="flex-direction:column;align-items:flex-start;gap:6px;">
      <div style="display:flex;align-items:center;gap:8px;flex-wrap:wrap;">
        <div class="tgl-group">
          <button class="tgl active" id="tgl-ollama"    onclick="setProvider('ollama')">Local (Ollama)</button>
          <button class="tgl"        id="tgl-anthropic" onclick="setProvider('anthropic')">Anthropic API</button>
          <button class="tgl"        id="tgl-hybrid"    onclick="setProvider('hybrid')" title="Ollama builds the knowledge base (better search), Claude answers (faster). Needs both.">⚡ Hybrid</button>
        </div>
        <div class="tgl-group">
          <button class="tgl active" id="tgl-whisper"   onclick="setTranscribeMode('whisper')">Whisper</button>
          <button class="tgl"        id="tgl-google"    onclick="setTranscribeMode('google')">Google</button>
          <button class="tgl"        id="tgl-deepgram"  onclick="setTranscribeMode('deepgram')">⚡ Deepgram</button>
        </div>
        <select id="whisper-model-select" title="Whisper model size (smaller = faster)">
          <option value="small" selected>small (fastest)</option>
          <option value="medium">medium (balanced)</option>
          <option value="large-v3-turbo">large-v3-turbo (best)</option>
        </select>
        <input id="deepgram-key-input" type="password" placeholder="Deepgram API key…"
               style="display:none;width:200px;" title="Get a free key at deepgram.com">
        <div class="tgl-group">
          <button class="tgl active" id="tgl-mic"    onclick="setAudioSource('mic')">🎤 Mic</button>
          <button class="tgl"        id="tgl-system" onclick="setAudioSource('system')">🖥️ System</button>
        </div>
      </div>
      <div id="ollama-controls" style="display:flex;align-items:center;gap:8px;flex-wrap:wrap;">
        <select id="model-select">
          <option value="gemma4:12b">gemma4:12b (best quality)</option>
          <option value="gemma3:4b">gemma3:4b (faster ~2s)</option>
          <option value="qwen2.5:3b">qwen2.5:3b (fastest ~1s)</option>
          <option value="gemma3:12b">gemma3:12b</option>
          <option value="mistral">mistral</option>
          <option value="llama3.1:8b">llama3.1:8b</option>
          <option value="custom">custom…</option>
        </select>
        <input id="model-custom" type="text" placeholder="model:tag" style="display:none;width:140px">
      </div>
      <div id="anthropic-controls" style="display:none;align-items:center;gap:8px;flex-wrap:wrap;">
        <select id="anthropic-model-select">
          <option value="claude-haiku-4-5-20251001">claude-haiku-4-5 (fastest)</option>
          <option value="claude-sonnet-4-6">claude-sonnet-4-6 (best quality)</option>
          <option value="claude-opus-4-8">claude-opus-4-8 (most capable)</option>
        </select>
        <input id="api-key-input" type="password" placeholder="sk-ant-api…" style="width:220px">
      </div>
      <div id="hybrid-controls" style="display:none;align-items:center;gap:8px;flex-wrap:wrap;">
        <span style="font-size:12px;opacity:.7;">Ollama KB model:</span>
        <select id="hybrid-ollama-model-select">
          <option value="gemma4:12b">gemma4:12b (recommended)</option>
          <option value="gemma3:12b">gemma3:12b</option>
          <option value="llama3.1:8b">llama3.1:8b</option>
        </select>
        <span style="font-size:12px;opacity:.7;">Claude answer model:</span>
        <select id="hybrid-anthropic-model-select">
          <option value="claude-haiku-4-5-20251001">claude-haiku-4-5 (fastest)</option>
          <option value="claude-sonnet-4-6">claude-sonnet-4-6 (best quality)</option>
        </select>
        <input id="hybrid-api-key-input" type="password" placeholder="sk-ant-api… (Anthropic key)" style="width:220px">
        <span style="font-size:11px;opacity:.6;">Ollama builds KB (better search) · Claude answers (faster streaming)</span>
      </div>
    </div>
    <div class="step-indicator ok" id="step1-ind">✓</div>
  </div>

  <div class="step">
    <div class="step-badge">2</div>
    <div class="step-label">Document</div>
    <div class="step-content">
      <span class="doc-name" id="doc-name">No file loaded</span>
      <input type="file" id="file-input" accept=".docx,.pdf">
      <button class="btn btn-blue" onclick="document.getElementById('file-input').click()">Browse…</button>
    </div>
    <div class="step-indicator" id="step2-ind"></div>
  </div>

  <div class="step">
    <div class="step-badge">3</div>
    <div class="step-label">Language</div>
    <div class="step-content">
      <select id="lang-select">
        <option value="es-ES" data-name="Spanish">Spanish</option>
        <option value="en-US" data-name="English">English</option>
        <option value="fr-FR" data-name="French">French</option>
        <option value="de-DE" data-name="German">German</option>
        <option value="it-IT" data-name="Italian">Italian</option>
        <option value="pt-BR" data-name="Portuguese">Portuguese</option>
      </select>
    </div>
    <div class="step-indicator ok" id="step3-ind">✓</div>
  </div>

  <button id="start-btn" class="btn btn-green" disabled onclick="toggleListening()">▶&nbsp;&nbsp;Start Listening</button>
  <button id="answer-now-btn" class="btn" style="display:none;background:var(--blue);color:#fff;" onclick="answerNow()" title="Answer now — works for both direct questions and after a clarification exchange">💬 Answer now</button>
  <hr style="margin-top:16px;">

  <div class="status-bar">
    <span class="status-dot" id="status-dot">●</span>
    <span class="status-text" id="status-text">Load a document to begin</span>
    <span id="mode-tag" class="tag" style="display:none"></span>
    <span style="margin-left:auto;" class="vol-bar-wrap">
      <span id="vol-label" style="font-size:11px;color:var(--dim);">mic</span>
      <span class="vol-bar-track"><span class="vol-bar-fill" id="vol-bar"></span></span>
    </span>
  </div>

  <div class="panel-header">
    <span class="section-label" style="margin:0">LIVE TRANSCRIPT</span>
    <button class="btn btn-dim" style="font-size:11px;padding:3px 8px;" onclick="clearTranscript()">Clear</button>
  </div>
  <div id="transcript-box"></div>

  <div class="answer-section">
    <div class="answer-header">
      <span class="section-label" style="margin:0">ANSWER</span>
      <span id="question-label"></span>
    </div>
    <div id="answer-card">
      <div id="answer-text">Answers will appear here when a question is detected…</div>
    </div>
  </div>
</div>

<script>
let isListening=false, audioContext=null, mediaStream=null, micStream=null, scriptProcessor=null, analyser=null;
let transcriptBuffer=[], answering=false;
let audioSamples=[], chunkTimer=null, speechActive=false, silenceSamples=0;
let transcribeMode='whisper', currentProvider='ollama', audioSourceMode='mic';
let _pendingAnswerTimer=null, _holdMode=false, _pendingCountdownInterval=null;
// Voice trigger — say any of these words to fire the answer hands-free
const VOICE_TRIGGERS = ['responde','contesta','contestar','answer','dale','ya'];
function _checkVoiceTrigger(text){
  const t = text.toLowerCase().replace(/[^a-záéíóúüñ\s]/g,'').trim();
  const words = t.split(/\s+/);
  // Match if the utterance IS (or starts/ends with) a trigger word — avoids false positives
  // e.g. "responde" → fire; "como responde" → fire; "corresponde" → no (not a word boundary match)
  return VOICE_TRIGGERS.some(trigger =>
    words[0] === trigger || words[words.length-1] === trigger ||
    (words.length <= 3 && words.includes(trigger))
  );
}
const SAMPLE_RATE=16000, SILENCE_RMS=0.002, SILENCE_END_S=0.5, MAX_UTTERANCE_S=5;

// Downsample from any source rate to 16kHz (linear interpolation).
// This is needed because getDisplayMedia on Windows delivers audio at the
// system native rate (44100 or 48000 Hz) and Chrome does NOT reliably
// resample display-capture audio when AudioContext is forced to 16 kHz.
function resampleTo16k(samples, fromRate) {
  if (fromRate === SAMPLE_RATE) return samples;
  const ratio = fromRate / SAMPLE_RATE;
  const outLen = Math.floor(samples.length / ratio);
  const out = new Float32Array(outLen);
  for (let i = 0; i < outLen; i++) {
    const src = i * ratio;
    const lo  = Math.floor(src);
    const hi  = Math.min(lo + 1, samples.length - 1);
    const frac = src - lo;
    out[i] = samples[lo] * (1 - frac) + samples[hi] * frac;
  }
  return out;
}

document.getElementById('model-select').addEventListener('change', () => {
  document.getElementById('model-custom').style.display =
    document.getElementById('model-select').value==='custom' ? '' : 'none';
});
function setProvider(p) {
  currentProvider = p;
  document.getElementById('tgl-ollama').classList.toggle('active', p==='ollama');
  document.getElementById('tgl-anthropic').classList.toggle('active', p==='anthropic');
  document.getElementById('tgl-hybrid').classList.toggle('active', p==='hybrid');
  document.getElementById('ollama-controls').style.display    = p==='ollama'    ? 'flex' : 'none';
  document.getElementById('anthropic-controls').style.display = p==='anthropic' ? 'flex' : 'none';
  document.getElementById('hybrid-controls').style.display    = p==='hybrid'    ? 'flex' : 'none';
}
function getModel() {
  if (currentProvider==='anthropic') return document.getElementById('anthropic-model-select').value;
  if (currentProvider==='hybrid')    return document.getElementById('hybrid-anthropic-model-select').value;
  const v=document.getElementById('model-select').value;
  return v==='custom' ? document.getElementById('model-custom').value.trim() : v;
}
function getApiKey() {
  if (currentProvider==='hybrid') return document.getElementById('hybrid-api-key-input').value.trim();
  return document.getElementById('api-key-input').value.trim();
}
function getOllamaModel() {
  if (currentProvider==='hybrid') return document.getElementById('hybrid-ollama-model-select').value;
  if (currentProvider==='ollama') {
    const v=document.getElementById('model-select').value;
    return v==='custom' ? document.getElementById('model-custom').value.trim() : v;
  }
  return '';
}
function setTranscribeMode(m) {
  transcribeMode = m;
  document.getElementById('tgl-whisper').classList.toggle('active', m === 'whisper');
  document.getElementById('tgl-google').classList.toggle('active', m === 'google');
  document.getElementById('tgl-deepgram').classList.toggle('active', m === 'deepgram');
  document.getElementById('whisper-model-select').style.display = m === 'whisper' ? '' : 'none';
  document.getElementById('deepgram-key-input').style.display  = m === 'deepgram' ? '' : 'none';
  if (isListening) {
    const t = document.getElementById('mode-tag');
    t.textContent = m === 'whisper' ? 'Whisper' : m === 'deepgram' ? '⚡ Deepgram' : 'Google Speech';
  }
}
function setAudioSource(m) {
  audioSourceMode = m;
  document.getElementById('tgl-mic').classList.toggle('active', m === 'mic');
  document.getElementById('tgl-system').classList.toggle('active', m === 'system');
  document.getElementById('vol-label').textContent = m === 'system' ? 'sys' : 'mic';
}

document.getElementById('file-input').addEventListener('change', async (e) => {
  const file=e.target.files[0]; if(!file) return;
  setStatus('Loading document…','var(--yellow)');
  const fd=new FormData();
  fd.append('file',file);
  fd.append('model',getModel());
  fd.append('provider',currentProvider);
  fd.append('api_key',getApiKey());
  const ollamaModel = getOllamaModel();
  if (ollamaModel) fd.append('ollama_model', ollamaModel);
  if (currentProvider==='anthropic') fd.append('anthropic_model', document.getElementById('anthropic-model-select').value);
  if (currentProvider==='hybrid')    fd.append('anthropic_model', document.getElementById('hybrid-anthropic-model-select').value);
  try {
    const data=await fetch('/load-document',{method:'POST',body:fd}).then(r=>r.json());
    if(data.error){setStatus('Error: '+data.error,'var(--red)');return;}
    const display=data.name.length>32?data.name.slice(0,29)+'…':data.name;
    document.getElementById('doc-name').textContent=`${display}  (${data.chars.toLocaleString()} chars)`;
    document.getElementById('doc-name').classList.add('loaded');
    setStep2Pending();
    const procModel = getModel();
    setStatus(`Processing document with ${procModel}… (this may take 1–2 min, please wait)`, 'var(--yellow)');
    pollProcessing();
  } catch(e){ setStatus('Upload error: '+e.message,'var(--red)'); }
});

function pollProcessing() {
  const poll=setInterval(async () => {
    try {
      const data=await fetch('/process-status').then(r=>r.json());
      if(data.error) {
        clearInterval(poll);
        // Processing failed — still allow questions via raw keyword retrieval fallback
        setStatus('Processing error: ' + data.error.slice(0, 80) + ' — using keyword fallback', 'var(--yellow)');
        setStep2Done();
      } else if(data.done) {
        clearInterval(poll);
        setStep2Done();
        if(!isListening) setStatus('Knowledge base ready — you can start listening', 'var(--green)');
      }
    } catch(_){}
  }, 2000);
}

function setStep2Pending() {
  const i=document.getElementById('step2-ind');
  i.textContent='⏳'; i.className='step-indicator'; i.style.color='var(--yellow)';
  document.getElementById('start-btn').disabled=true;
}
function setStep2Done() {
  const i=document.getElementById('step2-ind');
  i.textContent='✓'; i.className='step-indicator ok'; i.style.color='';
  updateStartBtn();
}
function updateStartBtn() {
  const ok=document.getElementById('step2-ind').classList.contains('ok');
  const btn=document.getElementById('start-btn');
  btn.disabled=!ok;
  if(ok&&!isListening){btn.className='btn btn-green';btn.textContent='▶  Start Listening';}
}

function toggleListening(){isListening?stopListening():startListening();}

// ── Deepgram WebSocket streaming ──────────────────────────────────────────────
let dgSocket=null, dgInterim='', dgFinalBuffer='', dgAnswerTimer=null;

let dgAudioQueue = [];  // buffer audio while WebSocket is still connecting

function dgConnect(apiKey, lang){
  dgAudioQueue = [];
  const langCode = lang.split('-')[0];  // 'es-ES' → 'es'
  // nova-2 is Deepgram's most accurate general model; boost thesis-specific terms
  const keywords = [
    'SD-WAN','Starlink','VPN','LTE','VSAT','IPSec','IKEv2','AES','TLS',
    'SIEM','SASE','Peplink','SpeedFusion','ISO','PHVA','MiPymes',
    'ciberseguridad','conectividad','arquitectura','resiliencia'
  ].map(k=>`keywords=${encodeURIComponent(k+':3')}`).join('&');
  const url = `wss://api.deepgram.com/v1/listen?model=nova-2-general`
    + `&language=${langCode}&smart_format=true&punctuate=true`
    + `&interim_results=true&endpointing=300&utterance_end_ms=1000`
    + `&encoding=linear16&sample_rate=${SAMPLE_RATE}&channels=1&${keywords}`;
  setStatus('🔌 Connecting to Deepgram…', 'var(--yellow)');
  dgSocket = new WebSocket(url, ['token', apiKey]);
  dgSocket.binaryType = 'arraybuffer';
  dgSocket.onopen = () => {
    console.log('[Deepgram] connected');
    setStatus('🎙  Listening…', 'var(--green)');
    dgInterim = '';
    // Flush any audio that arrived while we were connecting
    if (dgAudioQueue.length > 0) {
      console.log(`[Deepgram] flushing ${dgAudioQueue.length} queued buffers`);
      for (const buf of dgAudioQueue) dgSocket.send(buf);
      dgAudioQueue = [];
    }
  };
  dgSocket.onclose = (e) => {
    console.log('[Deepgram] closed', e.code, e.reason);
    dgSocket = null;
    dgAudioQueue = [];
    const msg = {
      1000: 'Deepgram session ended normally',
      1006: 'Deepgram connection lost — check internet',
      4002: 'Deepgram auth failed — check API key',
      4003: 'Deepgram quota exceeded',
    }[e.code] || `Deepgram closed (${e.code})`;
    setStatus('⚠️ ' + msg, 'var(--red)');
  };
  dgSocket.onerror = (e) => {
    console.error('[Deepgram] error', e);
    setStatus('Deepgram error — check API key and internet', 'var(--red)');
  };
  dgSocket.onmessage = (evt) => {
    try {
      const msg = JSON.parse(evt.data);
      const alt = msg?.channel?.alternatives?.[0];
      if(!alt) return;
      const text = (alt.transcript||'').trim();
      if(!text) return;
      if(msg.is_final){
        dgInterim = '';
        // Accumulate finals — full question may span multiple is_final events
        if(_checkVoiceTrigger(text)){
          // Trigger word detected — fire answer, don't add to buffer
          dgFinalBuffer = '';
          if(dgAnswerTimer){clearTimeout(dgAnswerTimer);dgAnswerTimer=null;}
          appendTranscript('▶ ' + text);
          updateAnswerNowBtn();
          if(!answering) fireAnswer();
          return;
        }
        dgFinalBuffer = (dgFinalBuffer + ' ' + text).trim();
        appendTranscript(text);
        transcriptBuffer.push(text);
        if(transcriptBuffer.length>12) transcriptBuffer.shift();
        updateAnswerNowBtn();
        // Smarter debounce:
        //   '?' present → 1.2s (question clearly finished)
        //   ends with tag-word (verdad, cierto, no?) → 1.8s
        //   otherwise → 3s (still building context before the actual question)
        const buf = dgFinalBuffer;
        const hasQ = buf.includes('?');
        const endsWithTag = /\b(verdad|correcto|cierto|no|eh)\s*[,.]?\s*$/.test(buf.toLowerCase());
        const debounceMs = hasQ ? 1200 : endsWithTag ? 1800 : 3000;
        if(dgAnswerTimer) clearTimeout(dgAnswerTimer);
        dgAnswerTimer = setTimeout(() => {
          const fullQuestion = dgFinalBuffer;
          dgFinalBuffer = '';
          dgAnswerTimer = null;
          updateAnswerNowBtn();
          // No auto-answer — user presses "Answer now" when ready (handles clarifications naturally)
        }, debounceMs);
      } else {
        // Show interim result live in status bar
        dgInterim = text;
        setStatus('🎙  ' + (dgFinalBuffer ? dgFinalBuffer + ' ' : '') + text.slice(-80), 'var(--green)');
      }
    } catch(_){}
  };
}

function dgSendAudio(samples){
  // samples is already resampled to 16kHz Float32Array
  if(!dgSocket) return;
  // Convert float32 → int16 PCM (Deepgram expects linear16)
  const buf = new Int16Array(samples.length);
  for(let i=0;i<samples.length;i++){
    const s=Math.max(-1,Math.min(1,samples[i]));
    buf[i] = s < 0 ? s * 0x8000 : s * 0x7FFF;
  }
  if(dgSocket.readyState === WebSocket.CONNECTING){
    // WebSocket not yet open — queue the buffer and flush on open
    dgAudioQueue.push(buf.buffer.slice(0));  // slice to own the buffer
    // Don't let the queue grow unbounded (keep last ~3s)
    if(dgAudioQueue.length > 8) dgAudioQueue.shift();
  } else if(dgSocket.readyState === WebSocket.OPEN){
    dgSocket.send(buf.buffer);
  }
}

function dgClose(){
  if(dgAnswerTimer){clearTimeout(dgAnswerTimer);dgAnswerTimer=null;}
  if(dgSocket){
    try{ dgSocket.send(JSON.stringify({type:'CloseStream'})); } catch(_){}
    dgSocket.close(); dgSocket=null;
  }
  dgInterim=''; dgFinalBuffer='';
  updateAnswerNowBtn();
}

function updateAnswerNowBtn(){
  const btn = document.getElementById('answer-now-btn');
  if (!btn) return;
  btn.style.display = (isListening && transcriptBuffer.length > 0) ? '' : 'none';
}

function answerNow(){
  // Manually trigger answer — cancel pending timers and answer immediately
  if(dgAnswerTimer){clearTimeout(dgAnswerTimer);dgAnswerTimer=null;}
  resetPendingAnswer();
  _holdMode = false;
  dgFinalBuffer = '';
  updateAnswerNowBtn();
  fireAnswer();
}

// Try to find Windows "Stereo Mix" / "What U Hear" loopback device.
// These appear as audio INPUT devices and work on any driver that supports them,
// unlike getDisplayMedia system audio which fails on many Windows audio drivers.
async function _findStereoMix() {
  try {
    // Must request mic permission first so labels are revealed
    await navigator.mediaDevices.getUserMedia({audio:true, video:false}).then(s=>s.getTracks().forEach(t=>t.stop()));
    const devices = await navigator.mediaDevices.enumerateDevices();
    const loopbackNames = ['stereo mix','what u hear','wave out mix','mixed output','sum','loopback',
                            'cable output','vb-audio','voicemeeter','virtual audio','blackhole','soundflower'];
    const found = devices.find(d =>
      d.kind === 'audioinput' &&
      loopbackNames.some(n => d.label.toLowerCase().includes(n))
    );
    return found || null;
  } catch(_) { return null; }
}

let displayStream = null;

async function startListening() {
  console.log('[Start] mode=' + audioSourceMode + ' transcribe=' + transcribeMode);
  // AudioContext MUST be created synchronously here, inside the click-handler gesture.
  // Any await (getDisplayMedia dialog, getUserMedia) expires the gesture — Chrome then
  // creates AudioContexts in suspended state and onaudioprocess delivers only zeros.
  audioContext = new AudioContext();
  console.log('[Start] AudioContext created, state=' + audioContext.state);

  try {
    if (audioSourceMode === 'system') {
      // Strategy 1: getDisplayMedia — captures ALL computer audio (WASAPI loopback on the
      // default render endpoint). Works regardless of output device (speakers or earphones).
      // User picks "Entire Screen" in the dialog and ticks "Share system audio".
      console.log('[Audio] Requesting getDisplayMedia for system audio…');
      const disp = await navigator.mediaDevices.getDisplayMedia({
        audio: {echoCancellation:false, noiseSuppression:false, autoGainControl:false},
        video: true,
        systemAudio: 'include'
      });
      displayStream = disp;
      const audioTrack = disp.getAudioTracks()[0];
      if (!audioTrack) {
        setStatus('No audio track — tick "Share system audio" in the Chrome dialog', 'var(--red)');
        disp.getTracks().forEach(t => t.stop());
        audioContext.close(); audioContext = null;
        return;
      }
      // Do NOT check audioTrack.muted here — Chrome reports system audio tracks as
      // temporarily muted during initialization on Windows; the track unmutes once
      // the WASAPI loopback opens. Checking muted here causes a false early exit.
      console.log('[Audio] getDisplayMedia OK, audioTrack:', audioTrack.label, 'muted:', audioTrack.muted, 'readyState:', audioTrack.readyState);
      mediaStream = new MediaStream([audioTrack]);
      audioTrack.onmute   = () => console.warn('[Audio] system track muted');
      audioTrack.onunmute = () => console.log('[Audio] system track unmuted');
      audioTrack.onended  = () => { if (isListening) stopListening(); };

      // Also capture mic so the user's own voice (clarifications) is transcribed too.
      // This is a best-effort add-on — failure here does NOT abort the session.
      try {
        micStream = await navigator.mediaDevices.getUserMedia({
          audio: { echoCancellation: true, noiseSuppression: true, autoGainControl: true },
          video: false
        });
        console.log('[Audio] Mic also captured — both voices will be transcribed');
      } catch(_) {
        micStream = null;
        console.warn('[Audio] Mic capture failed — system audio only (your voice won\'t be transcribed)');
      }
    } else {
      mediaStream = await navigator.mediaDevices.getUserMedia({audio:true, video:false});
    }
  } catch(e) {
    console.error('[Start] capture failed:', e);
    setStatus((audioSourceMode==='system' ? 'System audio capture' : 'Microphone') + ' failed: ' + e.message, 'var(--red)');
    if(audioContext){ audioContext.close(); audioContext = null; }
    return;
  }

  // Connect Deepgram WebSocket if in deepgram mode
  if (transcribeMode === 'deepgram') {
    const dgKey = document.getElementById('deepgram-key-input').value.trim();
    if (!dgKey) {
      setStatus('Enter your Deepgram API key first', 'var(--red)');
      mediaStream.getTracks().forEach(t => t.stop()); mediaStream = null;
      if (displayStream) { displayStream.getTracks().forEach(t=>t.stop()); displayStream=null; }
      audioContext.close(); audioContext = null;
      return;
    }
    const lang = document.getElementById('lang-select').value;
    dgConnect(dgKey, lang);
  }

  const nativeRate = audioContext.sampleRate;
  console.log(`[Audio] ctx=${nativeRate}Hz`);
  const captureLabel = audioSourceMode === 'system'
    ? (micStream ? 'System + Mic' : 'System audio')
    : 'Mic';
  setStatus(`🎙  Listening… (${captureLabel})`, 'var(--green)');

  const source = audioContext.createMediaStreamSource(mediaStream);
  analyser = audioContext.createAnalyser(); analyser.fftSize = 256; source.connect(analyser);

  const bufSize = Math.pow(2, Math.round(Math.log2(nativeRate * 0.256)));
  scriptProcessor = audioContext.createScriptProcessor(bufSize, 1, 1);
  source.connect(scriptProcessor);

  // In system mode, also mix in mic so both voices get transcribed
  if (micStream) {
    const micSource = audioContext.createMediaStreamSource(micStream);
    micSource.connect(scriptProcessor);
  }

  // Mute system audio to prevent feedback; ScriptProcessor must reach destination
  // or onaudioprocess won't fire.
  const silentOut = audioContext.createGain();
  silentOut.gain.value = audioSourceMode === 'system' ? 0 : 1;
  scriptProcessor.connect(silentOut);
  silentOut.connect(audioContext.destination);

  let _rmsLogCounter = 0;
  scriptProcessor.onaudioprocess=(e)=>{
    const raw = e.inputBuffer.getChannelData(0);
    // Resample from native rate down to 16kHz
    const samples = resampleTo16k(raw, nativeRate);
    const rms=Math.sqrt(samples.reduce((s,x)=>s+x*x,0)/samples.length);
    // Log first callback so user can check console for raw values
    if(_rmsLogCounter === 0){
      console.log(`[Audio] first callback: bufLen=${raw.length} rms=${rms.toFixed(5)} max=${Math.max(...raw.slice(0,16).map(Math.abs)).toFixed(5)} nCh=${e.inputBuffer.numberOfChannels}`);
    }
    if(++_rmsLogCounter % 20 === 0){
      console.log(`[Audio] rms=${rms.toFixed(5)}`);
      const bar = rms > 0.001 ? '▓'.repeat(Math.min(10,Math.round(rms*200))) : '';
      if(rms > 0.001){
        setStatus(`🎙  Listening… ${bar}`, 'var(--green)');
      } else if(_rmsLogCounter > 16 && audioSourceMode === 'system'){
        const trk = mediaStream && mediaStream.getAudioTracks()[0];
        const trkState = trk ? `track=${trk.readyState} muted=${trk.muted}` : 'no track';
        console.warn('[Audio] silence detected:', trkState);
        setStatus('⚠️ No audio — make sure you picked "Entire Screen" (not a window) and ticked "Also share system audio" in the Chrome dialog', 'var(--yellow)');
      }
    }
    if(transcribeMode==='deepgram'){
      dgSendAudio(samples);
      return;
    }
    // Always accumulate all samples (system audio bypasses VAD; mic uses it for early cuts)
    audioSamples.push(...samples);
    if(audioSourceMode!=='system'){
      // Mic: trigger early send on silence after speech
      if(rms>SILENCE_RMS){speechActive=true;silenceSamples=0;}
      else if(speechActive){
        silenceSamples+=samples.length;
        if(silenceSamples/SAMPLE_RATE>=SILENCE_END_S){
          speechActive=false; silenceSamples=0;
          if(chunkTimer){clearTimeout(chunkTimer);chunkTimer=null;} sendChunk();
        }
      }
    }
  };
  isListening=true; audioSamples=[]; speechActive=false; silenceSamples=0;
  const btn=document.getElementById('start-btn');
  btn.className='btn btn-red'; btn.textContent='⏹  Stop Listening'; btn.disabled=false;
  const tag=document.getElementById('mode-tag');
  tag.textContent=transcribeMode==='whisper' ? 'Whisper' : transcribeMode==='deepgram' ? '⚡ Deepgram' : 'Google Speech'; tag.style.display='';
  setStatus('🎙  Listening…','var(--green)'); updateVolume();
  function resetCap(){
    if(chunkTimer)clearTimeout(chunkTimer);
    chunkTimer=setTimeout(()=>{if(audioSamples.length>0)sendChunk();if(isListening)resetCap();},MAX_UTTERANCE_S*1000);
  }
  resetCap();
}

function stopListening(){
  isListening=false; clearTimeout(chunkTimer); chunkTimer=null;
  resetPendingAnswer(); _holdMode=false;
  dgClose();
  if(scriptProcessor){scriptProcessor.disconnect();scriptProcessor=null;}
  if(analyser){analyser.disconnect();analyser=null;}
  if(audioContext){audioContext.close();audioContext=null;}
  if(mediaStream){mediaStream.getTracks().forEach(t=>t.stop());mediaStream=null;}
  if(micStream){micStream.getTracks().forEach(t=>t.stop());micStream=null;}
  if(displayStream){displayStream.getTracks().forEach(t=>t.stop());displayStream=null;}
  audioSamples=[]; document.getElementById('vol-bar').style.width='0%';
  document.getElementById('mode-tag').style.display='none';
  document.getElementById('start-btn').className='btn btn-green';
  document.getElementById('start-btn').textContent='▶  Start Listening';
  document.getElementById('answer-now-btn').style.display='none';
  setStatus('Stopped','var(--dim)');
}

function updateVolume(){
  if(!isListening||!analyser)return;
  const buf=new Uint8Array(analyser.frequencyBinCount); analyser.getByteFrequencyData(buf);
  document.getElementById('vol-bar').style.width=Math.min(100,buf.reduce((a,b)=>a+b,0)/buf.length*2)+'%';
  requestAnimationFrame(updateVolume);
}

async function sendChunk(){
  if(!isListening||audioSamples.length===0)return;
  const samples=audioSamples.splice(0,audioSamples.length);
  const wav = encodeWav(samples, SAMPLE_RATE);
  const lang = document.getElementById('lang-select').value;
  try{
    const wmodel=document.getElementById('whisper-model-select').value;
    const data=await fetch(`/transcribe?lang=${lang}&mode=${transcribeMode}&wmodel=${wmodel}`,
      {method:'POST',headers:{'Content-Type':'audio/wav'},body:wav}).then(r=>r.json());
    if(data.error){setStatus('Transcription error: '+data.error,'var(--yellow)');return;}
    const text=(data.text||'').trim(); if(!text)return;
    if(_checkVoiceTrigger(text)){
      // User said the trigger word — fire answer without adding the trigger to the buffer
      appendTranscript('▶ ' + text);
      updateAnswerNowBtn();
      if(!answering) fireAnswer();
    } else {
      appendTranscript(text); transcriptBuffer.push(text);
      if(transcriptBuffer.length>12)transcriptBuffer.shift();
      updateAnswerNowBtn();
    }
  }catch(e){console.warn('Transcribe:',e.message);}
}

function encodeWav(samples,sr){
  const buf=new ArrayBuffer(44+samples.length*2),v=new DataView(buf);
  function ws(o,s){for(let i=0;i<s.length;i++)v.setUint8(o+i,s.charCodeAt(i));}
  ws(0,'RIFF');v.setUint32(4,36+samples.length*2,true);ws(8,'WAVE');ws(12,'fmt ');
  v.setUint32(16,16,true);v.setUint16(20,1,true);v.setUint16(22,1,true);
  v.setUint32(24,sr,true);v.setUint32(28,sr*2,true);v.setUint16(32,2,true);v.setUint16(34,16,true);
  ws(36,'data');v.setUint32(40,samples.length*2,true);
  for(let i=0;i<samples.length;i++){const s=Math.max(-1,Math.min(1,samples[i]));v.setInt16(44+i*2,s<0?s*0x8000:s*0x7FFF,true);}
  return buf;
}

function resetPendingAnswer() {
  if(_pendingAnswerTimer){ clearTimeout(_pendingAnswerTimer); _pendingAnswerTimer = null; }
  if(_pendingCountdownInterval){ clearInterval(_pendingCountdownInterval); _pendingCountdownInterval = null; }
}

function fireAnswer() {
  // Send the last 5 transcript chunks — captures a full clarification exchange if one happened
  const recentTurns = transcriptBuffer.slice(-5).join(' ').trim();
  if(recentTurns && !answering) answerQuestion(recentTurns);
}

function isQuestion(text){
  if(text.includes('?'))return true;
  const t=text.toLowerCase().trim();

  // Question words and command verbs in English, Spanish, French, German, Portuguese
  const sw=[
    // English question words + request verbs
    'what','how','why','when','where','who','which',
    'can','could','would','should','is','are','do','does','did','will','may','might','shall',
    'explain','tell','describe','show','define','summarize','list',
    // Spanish question words
    'qué','que','cómo','como','cuál','cual','cuáles','cuales',
    'cuándo','cuando','cuánto','cuanto','cuántos','cuantos',
    'dónde','donde','quién','quien','quiénes','quienes',
    'por qué','por que','puede','podría','debería','podrías',
    'es','son','hay','tiene','tienen','fue','fueron',
    // Spanish command/request verbs (these were missing — root cause of the bug)
    'explícame','explicame','explica','explicar',
    'dime','di','cuéntame','cuentame','cuéntanos',
    'háblame','hablame','háblanos','hablanos',
    'describe','describeme','describenos',
    'menciona','mencione','muéstrame','muestrame',
    'dame','danos','presenta','presente',
    // French
    "qu'est",'comment','pourquoi','quand','où','qui','quel','quelle',
    // German
    'was','wie','warum','wann','wo','wer','welche','welcher',
    // Portuguese
    'como','quando','onde','quem','qual','quais','me','nos'
  ];
  if(sw.some(w=>t.startsWith(w+' ')||t===w))return true;
  // Mid-sentence Spanish question words
  if(['qué','cuál','cuáles','cuándo','dónde','quién','cómo','que','cual'].some(w=>t.includes(' '+w+' ')))return true;
  // Fallback: treat as question if text is ≥4 words and starts with a verb-like word (covers remaining command forms)
  const words=t.split(/\s+/);
  if(words.length>=4){
    const first=words[0];
    // Ends in 'a' or 'e' → likely a conjugated Spanish verb (explícame, detalla, explica, etc.)
    if(/[ae]$/i.test(first)&&first.length>=4)return true;
  }
  return false;
}

async function answerQuestion(question){
  answering=true;
  const safety=setTimeout(()=>{answering=false;setStatus('🎙  Listening…','var(--green)');},20000);
  setStatus('Thinking…','var(--blue)');
  document.getElementById('question-label').textContent='Q: '+question;
  const sel=document.getElementById('lang-select');
  const langName=sel.options[sel.selectedIndex].dataset.name;
  const el=document.getElementById('answer-text'); el.textContent=''; el.className='has-answer';
  try{
    const resp=await fetch('/ask',{method:'POST',headers:{'Content-Type':'application/json'},
      body:JSON.stringify({question,language:langName,model:getModel(),provider:currentProvider,api_key:getApiKey()})});
    const reader=resp.body.getReader(); const dec=new TextDecoder(); let buf='';
    while(true){
      const {done,value}=await reader.read(); if(done)break;
      buf+=dec.decode(value,{stream:true});
      const lines=buf.split('\n'); buf=lines.pop();
      for(const line of lines){
        if(!line.startsWith('data: '))continue;
        const payload=line.slice(6).trim();
        if(payload==='[DONE]')break;
        try{
          const parsed=JSON.parse(payload);
          if(typeof parsed==='string') el.textContent+=parsed;
          else if(parsed.error){el.textContent='Error: '+parsed.error;el.className='';}
        }catch(_){}
      }
    }
    el.textContent=el.textContent.replace(/\*\*(.+?)\*\*/g,'$1').replace(/\*(.+?)\*/g,'$1').trim();
    if(!el.textContent) el.textContent='(no answer)';
  }catch(e){el.textContent='Failed: '+e.message;el.className='';}
  clearTimeout(safety); setStatus('🎙  Listening…','var(--green)'); answering=false;
}

function setStatus(msg,color){
  document.getElementById('status-text').textContent=msg;
  document.getElementById('status-dot').style.color=color;
  document.getElementById('status-text').style.color=color;
}
function appendTranscript(text){
  const box=document.getElementById('transcript-box'),span=document.createElement('span');
  span.textContent=text+' '; span.style.color='var(--text)';
  box.appendChild(span); box.scrollTop=box.scrollHeight;
}
function clearTranscript(){document.getElementById('transcript-box').innerHTML='';transcriptBuffer=[];}
</script>
</body>
</html>"""


# ── Entry point ────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    port = 5000
    url  = f"http://localhost:{port}"
    print(f"\n  Meeting Assistant v3 — streaming edition")
    print(f"  Running at {url}")
    print(f"  Make sure Ollama is running: ollama serve")
    print(f"  Default model: {DEFAULT_MODEL}")
    print(f"  Press Ctrl+C to stop\n")
    threading.Timer(1.2, lambda: webbrowser.open(url)).start()
    if _WHISPER_AVAILABLE:
        threading.Thread(target=_preload_whisper, daemon=True).start()
    app.run(host="127.0.0.1", port=port, debug=False, threaded=True)
